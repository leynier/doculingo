from pathlib import Path
from typing import Annotated

from docx import Document
from docx.text.paragraph import Paragraph
from typer import BadParameter, Exit, Option, echo

from ..common.errors import TranslationError
from ..common.translators import DEFAULT_MODEL, DEFAULT_RETRIES, Translator, get_translator
from .styler import copy_paragraph_style, copy_run_style


def _paragraph_text(para: Paragraph) -> str:
    """Return the text of a paragraph, preferring the single run when possible."""
    if len(para.runs) == 1:
        return para.runs[0].text
    return para.text


def _translatable_texts(paragraphs: list[Paragraph]) -> list[str]:
    """Return the non-blank texts of the paragraphs, in order."""
    return [text for text in map(_paragraph_text, paragraphs) if text.strip()]


def main(
    input_path: Annotated[
        Path,
        Option(
            "--input",
            "-i",
            exists=True,
            dir_okay=False,
            resolve_path=True,
            help="Input file path",
        ),
    ],
    output_path: Annotated[
        Path,
        Option(
            "--output",
            "-o",
            help="Output file path",
        ),
    ],
    language_source: Annotated[
        str,
        Option(
            "--language-source",
            "-s",
            help="Source language. For example: 'spanish'",
        ),
    ],
    language_target: Annotated[
        str,
        Option(
            "--language-target",
            "-t",
            help="Target language. For example: 'english'",
        ),
    ],
    translator: Annotated[
        Translator,
        Option(
            help="Translator",
            case_sensitive=False,
        ),
    ] = Translator.openai,
    model: Annotated[
        str,
        Option(
            "--model",
            help="Model used by the translator (only OpenAI for now).",
        ),
    ] = DEFAULT_MODEL,
    retries: Annotated[
        int,
        Option(
            "--retries",
            min=1,
            help="Number of retries after the first failed translation attempt.",
        ),
    ] = DEFAULT_RETRIES,
) -> None:
    if input_path.suffix.lower() != ".docx":
        raise BadParameter(
            f"Input file must be a .docx file, got '{input_path.name}'",
            param_hint="--input",
        )
    translate = get_translator(translator, model=model, retries=retries)
    doc = Document(str(input_path))
    paragraphs = list(doc.paragraphs)
    texts = _translatable_texts(paragraphs)
    print(f"Total paragraphs: {len(paragraphs)}")
    print(f"Total texts: {len(texts)}")
    translated: list[str] = []
    try:
        step = max(1, len(texts) // 100)
        for start in range(0, len(texts), step):
            batch = texts[start : start + step]
            result = translate(batch, language_source, language_target)
            if len(result) != len(batch):
                raise TranslationError(
                    f"The translator returned {len(result)} texts for a batch of {len(batch)} texts"
                )
            translated.extend(result)
    except (TranslationError, RuntimeError) as error:
        echo(f"Error: {error}", err=True)
        raise Exit(code=1) from error
    print(f"Total translated texts: {len(translated)}")
    translated_doc = Document()
    index = 0
    for para in paragraphs:
        if len(para.runs) == 1:
            text = para.runs[0].text
            if text.strip():
                text = translated[index]
                index += 1
            new_para = translated_doc.add_paragraph()
            copy_paragraph_style(para, new_para)
            new_run = new_para.add_run(text)
            copy_run_style(para.runs[0], new_run)
            continue
        text = para.text
        if text.strip():
            text = translated[index]
            index += 1
        new_para = translated_doc.add_paragraph(text)
        copy_paragraph_style(para, new_para)
    if output_path.suffix.lower() != ".docx":
        output_path = output_path.parent / (output_path.name + ".docx")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    translated_doc.save(str(output_path))
