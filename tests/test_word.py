from pathlib import Path

from docx import Document
from pytest import MonkeyPatch

from doculingo import word
from doculingo.common.translators import Translator


def _invoke(
    monkeypatch: MonkeyPatch,
    tmp_path: Path,
    translator_factory,
    *,
    input_texts: list[str] | None = None,
    paragraphs: list[list[str]] | None = None,
    **kwargs,
) -> tuple[Path, Path]:
    doc = Document()
    if paragraphs is not None:
        for runs in paragraphs:
            para = doc.add_paragraph()
            for run in runs:
                para.add_run(run)
    else:
        for text in input_texts or []:
            doc.add_paragraph(text)
    input_file = tmp_path / "input.docx"
    doc.save(str(input_file))
    output_file = tmp_path / "output.docx"
    if translator_factory is not None:
        monkeypatch.setattr(word, "get_translator", translator_factory)
    word.main(
        input_path=input_file,
        output_path=output_file,
        language_source="spanish",
        language_target="english",
        **kwargs,
    )
    return input_file, output_file


def test_segmentation(monkeypatch: MonkeyPatch, tmp_path: Path) -> None:
    texts = [f"text {i}" for i in range(250)]
    segments: list[list[str]] = []

    def fake_translate(items: list[str], ls: str, lt: str) -> list[str]:
        segments.append(list(items))
        return list(items)

    _, output_file = _invoke(
        monkeypatch, tmp_path, lambda *a, **k: fake_translate, input_texts=texts
    )

    step = len(texts) // 100
    expected = [texts[i : i + step] for i in range(0, len(texts), step)]
    assert segments == expected
    assert output_file.exists()


def test_options_forwarded_to_translator(monkeypatch: MonkeyPatch, tmp_path: Path) -> None:
    captured: dict[str, object] = {}

    def fake_get_translator(translator, model=None, retries=None):
        captured.update(translator=translator, model=model, retries=retries)

        def fake_translate(items: list[str], ls: str, lt: str) -> list[str]:
            return list(items)

        return fake_translate

    _invoke(
        monkeypatch,
        tmp_path,
        fake_get_translator,
        input_texts=["hola"],
        model="gpt-4o-mini",
        retries=3,
    )

    assert captured == {
        "translator": Translator.openai,
        "model": "gpt-4o-mini",
        "retries": 3,
    }


def test_no_translatable_paragraphs(monkeypatch: MonkeyPatch, tmp_path: Path) -> None:
    calls: list[list[str]] = []

    def fake_translate(items: list[str], ls: str, lt: str) -> list[str]:
        calls.append(list(items))
        return list(items)

    _, output_file = _invoke(
        monkeypatch,
        tmp_path,
        lambda *a, **k: fake_translate,
        input_texts=["", "   ", "\t\n"],
    )

    assert calls == []
    output_doc = Document(str(output_file))
    assert [para.text for para in output_doc.paragraphs] == ["", "   ", "\t\n"]


def test_whitespace_paragraphs_not_translated(monkeypatch: MonkeyPatch, tmp_path: Path) -> None:
    received: list[list[str]] = []

    def fake_translate(items: list[str], ls: str, lt: str) -> list[str]:
        received.append(list(items))
        return [text.upper() for text in items]

    _, output_file = _invoke(
        monkeypatch,
        tmp_path,
        lambda *a, **k: fake_translate,
        input_texts=["hola", "   ", "mundo"],
    )

    assert received == [["hola"], ["mundo"]]
    output_doc = Document(str(output_file))
    assert [para.text for para in output_doc.paragraphs] == ["HOLA", "   ", "MUNDO"]


def test_multi_run_paragraph_translated_as_whole_text(
    monkeypatch: MonkeyPatch, tmp_path: Path
) -> None:
    received: list[list[str]] = []

    def fake_translate(items: list[str], ls: str, lt: str) -> list[str]:
        received.append(list(items))
        return [text.upper() for text in items]

    _, output_file = _invoke(
        monkeypatch,
        tmp_path,
        lambda *a, **k: fake_translate,
        paragraphs=[["Hello ", "bold", " world"]],
    )

    assert received == [["Hello bold world"]]
    output_doc = Document(str(output_file))
    para = output_doc.paragraphs[0]
    assert para.text == "HELLO BOLD WORLD"
    assert len(para.runs) == 1
